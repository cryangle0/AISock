# -*- coding: utf-8 -*-
"""
爱花型袜版设计系统 实现方案 Markdown → Word 转换脚本
基于 E:\\angsa\\angsa_data\\AITools\\create_word(排版).py 排版规则
- 黑体二号/三号/四号标题、宋体小四正文、1.5倍行距、首行缩进 0.74cm
- 十章主体；前八章和"九、验收标准"前分页
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os


def set_run_font(run, font_name='宋体', font_size=None, bold=None):
    """设置字体（包括中文 East Asian 字形）"""
    run.font.name = font_name
    if font_size:
        run.font.size = Pt(font_size)
    if bold is not None:
        run.font.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.insert(0, rFonts)


def remove_heading_border(heading):
    """移除标题下方默认横线"""
    pPr = heading._element.get_or_add_pPr()
    for child in pPr:
        if child.tag.endswith('pBdr'):
            pPr.remove(child)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'none')
    bottom.set(qn('w:sz'), '0')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)


def parse_markdown(md_content):
    """解析 Markdown 为元素列表"""
    lines = md_content.split('\n')
    elements = []
    in_code_block = False
    code_content = []
    code_lang = ''
    table_buffer = []

    i = 0
    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lang = line.strip()[3:].strip()
                code_content = []
            else:
                in_code_block = False
                elements.append(('code', '\n'.join(code_content), code_lang))
            i += 1
            continue
        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # 表格
        stripped = line.strip()
        if stripped.startswith('|') and stripped.endswith('|'):
            if re.match(r'^\|[\s\-:|]+\|$', stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if cells:
                table_buffer.append(cells)
            i += 1
            continue
        elif table_buffer:
            elements.append(('table', table_buffer.copy()))
            table_buffer = []

        # 标题
        if re.match(r'^# [^#]', line):
            elements.append(('h1', line[2:].strip()))
            i += 1
            continue
        elif re.match(r'^## [^#]', line):
            elements.append(('h2', line[3:].strip()))
            i += 1
            continue
        elif re.match(r'^### [^#]', line):
            elements.append(('h3', line[4:].strip()))
            i += 1
            continue
        elif re.match(r'^#### [^#]', line):
            elements.append(('h4', line[5:].strip()))
            i += 1
            continue

        # 列表
        if line.strip().startswith('- '):
            elements.append(('list', line.strip()[2:]))
            i += 1
            continue
        if re.match(r'^\d+\. ', line.strip()):
            m = re.match(r'^\d+\. (.+)', line.strip())
            if m:
                elements.append(('numlist', m.group(1)))
            i += 1
            continue

        # 分隔线
        if line.strip() == '---':
            elements.append(('hr', ''))
            i += 1
            continue

        # 普通段落
        if line.strip():
            elements.append(('paragraph', line.strip()))

        i += 1

    if table_buffer:
        elements.append(('table', table_buffer))

    return elements


def create_word_document(md_file, output_file):
    """生成 Word 文档"""
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    doc = Document()

    # 默认字体：宋体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # A4 页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    elements = parse_markdown(md_content)
    print(f"解析到 {len(elements)} 个元素")
    type_count = {}
    for elem in elements:
        type_count[elem[0]] = type_count.get(elem[0], 0) + 1
    print(f"元素类型统计: {type_count}")

    # 七章章节起新页
    page_break_titles = ['一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、']

    for elem in elements:
        elem_type = elem[0]

        if elem_type == 'h1':
            # 文档标题：黑体二号(22pt) 居中 加粗
            h = doc.add_heading(elem[1], level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            remove_heading_border(h)
            for run in h.runs:
                set_run_font(run, '黑体', 22, bold=True)
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif elem_type == 'h2':
            title = elem[1]
            need_pb = any(title.startswith(t) for t in page_break_titles)
            if need_pb:
                doc.add_page_break()
            # 一级章节：黑体三号(16pt) 居中 加粗
            h = doc.add_heading(title, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in h.runs:
                set_run_font(run, '黑体', 16, bold=True)
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif elem_type == 'h3':
            # 题目：黑体四号(14pt) 左对齐 加粗
            h = doc.add_heading(elem[1], level=2)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in h.runs:
                set_run_font(run, '黑体', 14, bold=True)
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif elem_type == 'h4':
            # 题内子段：黑体小四(12pt) 左对齐 加粗
            h = doc.add_heading(elem[1], level=3)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in h.runs:
                set_run_font(run, '黑体', 12, bold=True)
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif elem_type == 'paragraph':
            text = elem[1]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing = 1.5

            # 处理 **粗体**
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_run_font(run, '宋体', 12, bold=True)
                elif part:
                    # 处理行内代码 `xxx`
                    inline_parts = re.split(r'(`[^`]+`)', part)
                    for ip in inline_parts:
                        if ip.startswith('`') and ip.endswith('`'):
                            run = p.add_run(ip[1:-1])
                            run.font.name = 'Consolas'
                            run.font.size = Pt(11)
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
                        elif ip:
                            run = p.add_run(ip)
                            set_run_font(run, '宋体', 12)

        elif elem_type == 'list':
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.line_spacing = 1.5
            text = elem[1]
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_run_font(run, '宋体', 12, bold=True)
                elif part:
                    run = p.add_run(part)
                    set_run_font(run, '宋体', 12)

        elif elem_type == 'numlist':
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(elem[1])
            set_run_font(run, '宋体', 12)

        elif elem_type == 'code':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            for line in elem[1].split('\n'):
                run = p.add_run(line + '\n')
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')

        elif elem_type == 'table':
            rows = elem[1]
            if rows:
                num_cols = max(len(row) for row in rows)
                table = doc.add_table(rows=len(rows), cols=num_cols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for ri, row_data in enumerate(rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < num_cols:
                            cell = table.rows[ri].cells[ci]
                            cell.text = ''
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run(cell_text)
                            set_run_font(run, '宋体', 10.5,
                                         bold=(ri == 0))
                doc.add_paragraph()

        elif elem_type == 'hr':
            # 横线由分页符承担，跳过
            pass

    doc.save(output_file)
    print(f'\nWord 文档已生成: {output_file}')


if __name__ == '__main__':
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_file = os.path.join(script_dir, '实现方案.md')
    output_file = os.path.join(
        script_dir, '爱花型袜版设计系统_实现方案.docx'
    )
    print(f"输入 MD: {md_file}")
    print(f"输出 DOCX: {output_file}\n")
    create_word_document(md_file, output_file)
